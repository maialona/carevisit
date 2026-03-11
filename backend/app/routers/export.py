"""PDF and Word export endpoints for visit records."""
from __future__ import annotations

import io
import uuid
import urllib.parse
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.database import get_db
from app.deps import get_current_user
from app.models.models import Organization, User, VisitRecord
from app.utils.html_to_docx import html_to_docx_paragraphs

router = APIRouter(prefix="/records", tags=["export"])


# ---------- helpers ----------

async def _get_record_context(
    record_id: uuid.UUID, db: AsyncSession, current_user: User
):
    result = await db.execute(select(VisitRecord).where(VisitRecord.id == record_id))
    record = result.scalar_one_or_none()
    if record is None:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="紀錄不存在")

    user_result = await db.execute(select(User).where(User.id == record.user_id))
    recorder = user_result.scalar_one_or_none()
    if not recorder or recorder.org_id != current_user.org_id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="權限不足")

    org_result = await db.execute(select(Organization).where(Organization.id == current_user.org_id))
    org = org_result.scalar_one_or_none()

    return record, recorder, org


def _strip_html(html: str) -> str:
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text("\n", strip=True)


def _visit_type_label(vt: str) -> str:
    return "家庭訪視" if vt == "home" else "電話訪視"


def _make_filename(case_name: str, visit_date: datetime, ext: str) -> str:
    date_str = visit_date.strftime("%Y%m%d")
    return f"訪視紀錄_{case_name}_{date_str}.{ext}"


# ---------- PDF ----------

# Register MicroHei font for better Chinese support
_FONT_PATH = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"
try:
    pdfmetrics.registerFont(TTFont("MicroHei", _FONT_PATH))
    _DEFAULT_FONT = "MicroHei"
except Exception:
    # Fallback to standard CID font if TTF not found
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    _DEFAULT_FONT = "STSong-Light"

_BASE_STYLE = ParagraphStyle(
    "ZhBase",
    fontName=_DEFAULT_FONT,
    fontSize=10,
    leading=14,
)


@router.get("/{record_id}/export/pdf")
async def export_pdf(
    record_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    record, recorder, org = await _get_record_context(record_id, db, current_user)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ZhTitle", parent=_BASE_STYLE,
        fontSize=16, alignment=1, spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "ZhHeading", parent=_BASE_STYLE,
        fontSize=12, spaceBefore=12, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "ZhBody", parent=_BASE_STYLE,
        fontSize=10, leading=16, spaceAfter=4,
    )

    elements = []

    # Header
    vt_label = _visit_type_label(record.visit_type.value)
    elements.append(Paragraph(f"{vt_label}紀錄表", title_style))
    elements.append(Spacer(1, 0.5 * cm))

    # Info table
    recorder_name = recorder.name if recorder else ""
    date_str = record.visit_date.strftime("%Y/%m/%d")
    info_data = [
        [
            Paragraph(f"個案姓名：{record.case_name}", _BASE_STYLE),
            Paragraph(f"訪視日期：{date_str}", _BASE_STYLE),
        ],
        [
            Paragraph(f"督導員：{recorder_name}", _BASE_STYLE),
            Paragraph(f"訪視類型：{vt_label}", _BASE_STYLE),
        ],
    ]
    info_table = Table(info_data, colWidths=[8 * cm, 8 * cm])
    info_table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(info_table)
    elements.append(Spacer(1, 0.5 * cm))

    # Content
    elements.append(Paragraph("訪視紀錄內容", heading_style))
    plain = _strip_html(record.refined_content or record.raw_input or "（無內容）")
    for line in plain.split("\n"):
        line = line.strip()
        if line:
            # Escape XML special chars for reportlab
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            elements.append(Paragraph(safe, body_style))

    elements.append(Spacer(1, 1.5 * cm))

    # Signature
    elements.append(Paragraph(
        f"督導員簽名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿＿＿＿＿",
        ParagraphStyle("Sig", parent=_BASE_STYLE, fontSize=10),
    ))

    doc.build(elements)
    buf.seek(0)

    filename = _make_filename(record.case_name, record.visit_date, "pdf")
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )


# ---------- DOCX ----------

@router.get("/{record_id}/export/docx")
async def export_docx(
    record_id: uuid.UUID,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    record, recorder, org = await _get_record_context(record_id, db, current_user)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    vt_label = _visit_type_label(record.visit_type.value)
    title = doc.add_heading(f"{vt_label}紀錄表", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info table
    recorder_name = recorder.name if recorder else ""
    date_str = record.visit_date.strftime("%Y/%m/%d")

    table = doc.add_table(rows=2, cols=2, style="Table Grid")
    table.cell(0, 0).text = f"個案姓名：{record.case_name}"
    table.cell(0, 1).text = f"訪視日期：{date_str}"
    table.cell(1, 0).text = f"督導員：{recorder_name}"
    table.cell(1, 1).text = f"訪視類型：{vt_label}"

    doc.add_paragraph()

    # Content heading
    doc.add_heading("訪視紀錄內容", level=2)

    # Parse HTML content
    content_html = record.refined_content or record.raw_input or ""
    if content_html.strip().startswith("<"):
        html_to_docx_paragraphs(doc, content_html)
    else:
        for line in content_html.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph()

    # Signature
    sig = doc.add_paragraph()
    sig_run = sig.add_run("督導員簽名：＿＿＿＿＿＿＿＿　　日期：＿＿＿＿＿＿＿＿")
    sig_run.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = _make_filename(record.case_name, record.visit_date, "docx")
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"},
    )
