"""Convert TipTap HTML to python-docx paragraphs."""
from __future__ import annotations

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt


def html_to_docx_paragraphs(doc: Document, html: str) -> None:
    """Parse TipTap HTML and write into a python-docx Document."""
    soup = BeautifulSoup(html, "html.parser")

    for element in soup.children:
        if isinstance(element, NavigableString):
            text = str(element).strip()
            if text:
                doc.add_paragraph(text)
        elif isinstance(element, Tag):
            _process_tag(doc, element)


def _process_tag(doc: Document, tag: Tag) -> None:
    name = tag.name

    if name in ("h1", "h2", "h3"):
        level = int(name[1])
        p = doc.add_heading(level=level)
        _add_inline_runs(p, tag)

    elif name == "ul":
        for li in tag.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, li)

    elif name == "ol":
        for li in tag.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, li)

    elif name == "p":
        p = doc.add_paragraph()
        _add_inline_runs(p, tag)

    elif name == "br":
        doc.add_paragraph()

    else:
        # Fallback: treat unknown tags as paragraphs
        text = tag.get_text(strip=True)
        if text:
            doc.add_paragraph(text)


def _add_inline_runs(paragraph, tag: Tag) -> None:
    """Walk inline children and create runs with bold/italic."""
    for child in tag.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
        elif isinstance(child, Tag):
            if child.name == "br":
                paragraph.add_run("\n")
            elif child.name in ("strong", "b"):
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ("em", "i"):
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "u":
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name in ("ul", "ol"):
                # Nested lists – flatten as indented text
                for li in child.find_all("li", recursive=False):
                    run = paragraph.add_run(f"\n  - {li.get_text()}")
                    run.font.size = Pt(10)
            else:
                # Recurse for nested inline tags
                _add_inline_runs(paragraph, child)
